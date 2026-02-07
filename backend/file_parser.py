import os
from PyPDF2 import PdfReader
from docx import Document

def extract_text_from_file(filepath):
    """
    Extract text content from various file types.
    Returns a dictionary with text content and metadata.
    """
    file_extension = os.path.splitext(filepath)[1].lower()

    try:
        if file_extension == '.pdf':
            return extract_text_from_pdf(filepath)
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(filepath)
        elif file_extension == '.txt':
            return extract_text_from_txt(filepath)
        else:
            return {
                'text': None,
                'error': f'Unsupported file type: {file_extension}',
                'pages': 0
            }
    except Exception as e:
        return {
            'text': None,
            'error': f'Error parsing file: {str(e)}',
            'pages': 0
        }

def extract_text_from_pdf(filepath):
    """Extract text from PDF file"""
    reader = PdfReader(filepath)
    text = ""

    for page in reader.pages:
        text += page.extract_text() + "\n\n"

    return {
        'text': text.strip(),
        'pages': len(reader.pages),
        'type': 'PDF'
    }

def extract_text_from_docx(filepath):
    """Extract text from Word document"""
    doc = Document(filepath)
    text = ""

    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    return {
        'text': text.strip(),
        'pages': len(doc.paragraphs),
        'type': 'Word Document'
    }

def extract_text_from_txt(filepath):
    """Extract text from plain text file"""
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()

    return {
        'text': text.strip(),
        'pages': 1,
        'type': 'Text File'
    }
